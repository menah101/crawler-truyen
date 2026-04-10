"""
docx_exporter.py — Xuất truyện đã viết lại ra file DOCX.

Cấu trúc giống hệt file mẫu vuong_phu_tam_bao.docx:
  - Normal style: Times New Roman 13pt
  - Lời chào → dòng trống → đoạn văn → dòng trống → ... → lời kết
  - KHÔNG có tiêu đề chương, KHÔNG dùng heading
  - align=LEFT, space_after=Pt(6) cho mỗi đoạn nội dung

Dùng khi DOCX_EXPORT_ENABLED = True trong config.py.
"""

import os
import re
import logging

try:
    from vi_validator import process_text as _vi_process_text
except ImportError:
    from crawler.vi_validator import process_text as _vi_process_text  # type: ignore

logger = logging.getLogger(__name__)


def _slugify(text: str) -> str:
    """Slug ASCII để đặt tên file."""
    VIET = {
        'à':'a','á':'a','ả':'a','ã':'a','ạ':'a',
        'ă':'a','ắ':'a','ằ':'a','ẳ':'a','ẵ':'a','ặ':'a',
        'â':'a','ấ':'a','ầ':'a','ẩ':'a','ẫ':'a','ậ':'a',
        'è':'e','é':'e','ẻ':'e','ẽ':'e','ẹ':'e',
        'ê':'e','ế':'e','ề':'e','ể':'e','ễ':'e','ệ':'e',
        'ì':'i','í':'i','ỉ':'i','ĩ':'i','ị':'i',
        'ò':'o','ó':'o','ỏ':'o','õ':'o','ọ':'o',
        'ô':'o','ố':'o','ồ':'o','ổ':'o','ỗ':'o','ộ':'o',
        'ơ':'o','ớ':'o','ờ':'o','ở':'o','ỡ':'o','ợ':'o',
        'ù':'u','ú':'u','ủ':'u','ũ':'u','ụ':'u',
        'ư':'u','ứ':'u','ừ':'u','ử':'u','ữ':'u','ự':'u',
        'ỳ':'y','ý':'y','ỷ':'y','ỹ':'y','ỵ':'y','đ':'d',
    }
    t = text.lower()
    t = ''.join(VIET.get(c, c) for c in t)
    t = re.sub(r'[^a-z0-9]+', '-', t)
    return t.strip('-')[:80]


def clean_text(text: str) -> str:
    """Làm sạch nội dung trước khi xuất DOCX: xóa tiêu đề chương, ký tự rác, thay từ nhạy cảm."""
    if not text:
        return text

    VIETNAMESE_CHARS = r'a-zA-ZÀ-ÁÂÃÈ-ÉÊÌ-ÍÒ-ÓÔÕÙ-ÚĂĐĨŨƠà-áâãè-éêì-íò-óôõù-úăđĩũơƯ-ưẠ-Ỹỳ-ỷỹ'
    VIETNAMESE_WORD  = f'[{VIETNAMESE_CHARS}0-9_]'

    # Xóa dòng tiêu đề chương (Chương 1, chương II, ...)
    text = re.sub(
        r'(?im)^[^\S\r\n]*chương\s+([IVXLCDM]+|\d+)(\s*[:\-–—.]?\s*.*)?\s*$',
        '', text,
    )
    # Xóa dòng "—HẾT—" và biến thể
    text = re.sub(
        r'(?im)^[^\S\r\n]*[-–—~*\[\](【】]?\s*H[EÉẾ][Tt]\s*[-–—~*\[\](【】]?\s*$',
        '', text,
    )
    # Xóa số đứng đầu dòng trước chữ hoa (VD: "1 Tôi" → "Tôi")
    text = re.sub(
        r'(?im)^\s*\d+\s+(?=[A-ZÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠƯẠ-Ỹ][a-zàáâãèéêìíòóôõùúăđĩũơưạ-ỹ])',
        '', text,
    )
    # Xóa dòng chỉ chứa dấu gạch (_____, --------, ════════)
    text = re.sub(r'(?im)^[^\S\r\n]*[-_─—═=]{3,}\s*$', '', text)
    # Xóa dòng "số.Chữ" (VD: "9.Sáng", "12.Hôm nay")
    text = re.sub(
        r'(?im)^[^\S\r\n]*\d+\.[A-ZÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠƯẠ-Ỹ].*$',
        '', text,
    )
    # Gộp nhiều dòng trống
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Xóa dấu chấm trong từ (VD: "T.i.ế.n" → "Tiến")
    def remove_dots_in_word(match):
        return match.group(1).replace('.', '')
    text = re.sub(rf'(\b[{VIETNAMESE_CHARS}]+\.[{VIETNAMESE_CHARS}\.]+\b)', remove_dots_in_word, text)

    # Xóa "***g"
    text = re.sub(r'\s*\*{{3}}g\s*', ' ', text)

    # Xóa Mathematical Alphanumeric Symbols và Emoji
    text = re.sub(r'[\U0001D400-\U0001D7FF]', '', text)
    text = re.sub(r'[\U0001F300-\U0001FAFF]', '', text)

    # Thay từ nhạy cảm
    replacement_map = {
        r'\bchó má\b':    'ngu ngốc',
        r'\bchết\b':      'chít',
        r'\bgiết\b':      'xử',
        r'\bchém\b':      'xém',
        r'\bmáu\b':       'huyết',
        r'\bmáu me\b':    'thảm khốc',
        r'\bmáu lạnh\b':  'tàn nhẫn',
        r'\buống máu\b':  'hút cạn sinh lực',
        r'\bngực\b':      'vòng một',
        r'blogger':       'bờ lóc gơ',
        r'vlog':          'vi lóc',
        r'audio':         'au đi ô',
    }
    for pattern, replacement in replacement_map.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # Tách câu bị dính liền: "nhẫnGió" → "nhẫn\nGió" (chữ thường + chữ HOA liền nhau)
    # Chỉ tách khi chữ trước KHÔNG phải tên riêng nước ngoài thường gặp
    _VIET_LOWER = r'[a-zàáảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ]'
    _VIET_UPPER = r'[A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴĐ]'
    text = re.sub(
        rf'({_VIET_LOWER})({_VIET_UPPER})',
        r'\1\n\2', text,
    )

    # Xóa ký tự đặc biệt không hợp lệ (giữ nguyên \n để tách đoạn)
    text = re.sub(rf'[^{VIETNAMESE_WORD}\s.,!?\'\"""\'\'…\-–—:;()\[\]{{}}\n]', '', text)
    # Gộp nhiều space trên cùng dòng (KHÔNG gộp \n)
    text = re.sub(r'[^\S\n]+', ' ', text)
    # Gộp nhiều \n liên tiếp thành 1
    text = re.sub(r'\n+', '\n', text).strip()

    # ── Kiểm tra & sửa chính tả phổ biến ─────────────────────────
    text = _fix_common_typos(text)

    return text


# ─────────────────────────────────────────────────────────────────
# Kiểm tra chính tả — sửa lỗi phổ biến từ AI rewrite tiếng Việt
# ─────────────────────────────────────────────────────────────────

# Từ sai → từ đúng (AI hay viết sai)
_TYPO_MAP = {
    # Lỗi AI hay gặp: thiếu dấu / sai dấu
    'nguoi': 'người',
    'khong': 'không',
    'nhung': 'nhưng',
    'duoc': 'được',
    'cung': 'cũng',
    'dang': 'đang',
    'nhin': 'nhìn',
    'muon': 'muốn',
    'biet': 'biết',
    'dieu': 'điều',
    'truoc': 'trước',
    'phai': 'phải',
    'chuyen': 'chuyện',
    'noi': 'nói',
    'tinh': 'tình',
    'nhu': 'như',
    'voi': 'với',
    'trong': 'trong',
    'anh': 'anh',
    # Lỗi phổ biến từ dịch máy
    'thế thế': 'thế',
    'rằng rằng': 'rằng',
    'mà mà': 'mà',
    'là là': 'là',
    'và và': 'và',
    'nhưng nhưng': 'nhưng',
    'của của': 'của',
    # Lỗi thiếu khoảng trắng sau dấu câu
}

# ── Lỗi OCR / ký tự bị merge / hỏng ──────────────────────────
# CHÚ Ý: Các lỗi thuần túy ký tự lặp (ngưười, nhữững, khôông, ươương, ...) đã
# được xử lý tự động bằng `vi_validator.process_text()` — KHÔNG cần liệt kê
# trong _OCR_MAP nữa. Map này chỉ giữ các lỗi ngữ cảnh mà validator không
# suy luận ra được (từ ghép đặc biệt, typo cố định hay gặp).
_OCR_MAP = {
    # Lỗi ngữ cảnh: AI sinh ra cụm từ đúng chính tả nhưng sai nghĩa
    # hoặc ghép 2 từ thành 1 âm tiết lạ (validator không đoán được).
    'vậy sánh': 'so sánh',
    'bị xúc Lưu': 'bị xúc phạm',
}

# ── Content filter workaround ──────────────────────────
# DeepSeek thay "chết" → "chít" để né content filter bạo lực.
# Phải thay lại vì "chít" không phải từ đúng trong ngữ cảnh tử vong.
# Chỉ áp dụng khi "chít" đi kèm từ context liên quan đến chết/hấp hối.
_DEATH_EUPHEMISM_PATTERNS = [
    # "cái chít" → "cái chết" (chắc chắn 100%)
    (r'\bcái\s+chít\b', 'cái chết'),
    # "bị chít" → "bị chết"
    (r'\bbị\s+chít\b', 'bị chết'),
    # "đã chít" → "đã chết"
    (r'\bđã\s+chít\b', 'đã chết'),
    # "chít oan" → "chết oan"
    (r'\bchít\s+oan\b', 'chết oan'),
    # "chít thảm" → "chết thảm"
    (r'\bchít\s+thảm\b', 'chết thảm'),
    # "chít rồi" → "chết rồi"
    (r'\bchít\s+rồi\b', 'chết rồi'),
    # "ai chít" → "ai chết"
    (r'\bai\s+chít\b', 'ai chết'),
    # "chít vì" → "chết vì"
    (r'\bchít\s+vì\b', 'chết vì'),
    # "tự chít" → "tự chết"
    (r'\btự\s+chít\b', 'tự chết'),
    # "sắp chít" → "sắp chết"
    (r'\bsắp\s+chít\b', 'sắp chết'),
    # "chít trong/tại/ở" → "chết trong/tại/ở"
    (r'\bchít\s+(trong|tại|ở|dưới|trên)\b', r'chết \1'),
    # "người chít" → "người chết"
    (r'\bngười\s+chít\b', 'người chết'),
    # "vụ chít" → "vụ chết"
    (r'\bvụ\s+chít\b', 'vụ chết'),
    # "cảnh chít" → "cảnh chết"
    (r'\bcảnh\s+chít\b', 'cảnh chết'),
]

# Từ tiếng Anh phổ biến trong truyện → tiếng Việt chuẩn
_ENGLISH_VIET_MAP = {
    'shock':       'sốc',
    'shocked':     'sốc',
    'shocking':    'sốc',
    'ok':          'ổn',
    'okay':        'được',
    'sorry':       'xin lỗi',
    'baby':        'em bé',
    'honey':       'cưng',
    'darling':     'cưng',
    'please':      'làm ơn',
    'thank':       'cảm ơn',
    'thanks':      'cảm ơn',
    'happy':       'vui vẻ',
    'bye':         'tạm biệt',
    'goodbye':     'tạm biệt',
    'hello':       'xin chào',
    'like':        'thích',
    'love':        'yêu',
    'cool':        'tuyệt',
    'nice':        'tốt',
    'hot':         'nóng',
    'sexy':        'quyến rũ',
    'stress':      'căng thẳng',
    'style':       'phong cách',
    'trend':       'xu hướng',
    # KHÔNG map 'so' → 'vậy': 'so' là từ Việt hợp lệ (so sánh, so với)
    'look':        'nhìn',
    'kiss':        'hôn',
    'sweet':       'ngọt ngào',
    'perfect':     'hoàn hảo',
    'romantic':    'lãng mạn',
    'crazy':       'điên',
    'single':      'độc thân',
    'couple':      'cặp đôi',
    'date':        'hẹn hò',
    'fans':        'người hâm mộ',
    'fan':         'người hâm mộ',
    'online':      'trực tuyến',
    'comeback':    'trở lại',
    'deadline':    'hạn chót',
    'feedback':    'phản hồi',
    'game':        'trò chơi',
    'idea':        'ý tưởng',
    'item':        'vật phẩm',
    'level':       'cấp độ',
    'list':        'danh sách',
    'comment':     'bình luận',
    'share':       'chia sẻ',
    'view':        'lượt xem',
    'check':       'kiểm tra',
    'update':      'cập nhật',
    'drama':       'bi kịch',
    'star':        'ngôi sao',
    'hot girl':    'gái xinh',
    'hot boy':     'trai đẹp',
    'full':        'đầy đủ',
    # Từ AI hay để nguyên tiếng Anh trong văn cảnh Việt
    'teaches':     'dạy',
    'slightly':    'hơi',
    'momento':     'khoảnh khắc',
    'voice':       'giọng',
    'moment':      'khoảnh khắc',
    'anyway':      'dù sao',
    'maybe':       'có lẽ',
    'really':      'thật sự',
    'actually':    'thực ra',
    'seriously':   'nghiêm túc',
    'important':   'quan trọng',
    'special':     'đặc biệt',
    'beautiful':   'xinh đẹp',
    'dangerous':   'nguy hiểm',
    'impossible':  'không thể',
    'surprise':    'bất ngờ',
    'feeling':     'cảm giác',
    'amazing':     'tuyệt vời',
    'terrible':    'kinh khủng',
    'wonderful':   'tuyệt vời',
    'horrible':    'khủng khiếp',
    'lucky':       'may mắn',
    'power':       'sức mạnh',
    'problem':     'vấn đề',
    'trouble':     'rắc rối',
    'secret':      'bí mật',
    'story':       'câu chuyện',
    'heart':       'trái tim',
    'chance':      'cơ hội',
    'control':     'kiểm soát',
    'memory':      'ký ức',
    'forever':     'mãi mãi',
    'future':      'tương lai',
    'together':    'cùng nhau',
    'trust':       'tin tưởng',
    'focus':       'tập trung',
}

# ── Phát hiện từ hỏng do AI rewriter ─────────────────────────
# Trong tiếng Việt, mỗi âm tiết chỉ có các tổ hợp nguyên âm hợp lệ.
# AI rewriter đôi khi sinh ra ký tự bị merge/hỏng tạo thành tổ hợp nguyên âm
# không tồn tại trong tiếng Việt.
_TONED_VOWELS = set(
    'àáảãạằắẳẵặầấẩẫậèéẻẽẹềếểễệìíỉĩịòóỏõọồốổỗộờớởỡợùúủũụừứửữựỳýỷỹỵ'
    'ÀÁẢÃẠẰẮẲẴẶẦẤẨẪẬÈÉẺẼẸỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌỒỐỔỖỘỜỚỞỠỢÙÚỦŨỤỪỨỬỮỰỲÝỶỸỴ'
)
_UNTONED_VOWELS = set('aăâeêioôơuưyAĂÂEÊIOÔƠUƯY')
_ALL_VOWELS = _TONED_VOWELS | _UNTONED_VOWELS

# Bảng strip dấu thanh về nguyên âm gốc
_STRIP_TONE = {}
_bases = [
    ('aàáảãạ', 'a'), ('ăằắẳẵặ', 'ă'), ('âầấẩẫậ', 'â'),
    ('eèéẻẽẹ', 'e'), ('êềếểễệ', 'ê'), ('iìíỉĩị', 'i'),
    ('oòóỏõọ', 'o'), ('ôồốổỗộ', 'ô'), ('ơờớởỡợ', 'ơ'),
    ('uùúủũụ', 'u'), ('ưừứửữự', 'ư'), ('yỳýỷỹỵ', 'y'),
]
for chars, base in _bases:
    for c in chars:
        _STRIP_TONE[c] = base
        _STRIP_TONE[c.upper()] = base.upper() if base != 'đ' else 'Đ'
# Nguyên âm không dấu thanh → giữ nguyên
for c in 'aăâeêioôơuưyAĂÂEÊIOÔƠUƯY':
    _STRIP_TONE[c] = c

# Tổ hợp nguyên âm HỢP LỆ trong tiếng Việt (sau khi strip dấu thanh, lowercase)
_VALID_VOWEL_CLUSTERS = {
    # Đôi
    'ai', 'ao', 'au', 'ay', 'âu', 'ây',
    'eo', 'êu',
    'ia', 'iê', 'iu',
    'oa', 'oă', 'oe', 'oi', 'oo', 'ôi', 'ơi',
    'ua', 'uâ', 'ue', 'uê', 'ui', 'uo', 'uô', 'uơ', 'uy',
    'ưa', 'ưi', 'ươ', 'ưu',
    'ya', 'yê',
    # Ba
    'oai', 'oay', 'oeo',
    'uây', 'uya', 'uyê', 'uyu',
    'ươi', 'ươu',
    'iêu', 'yêu',
    # Bốn
    'uyên',  # technically vowel+n but appears in clusters
}

# Ký tự không phải tiếng Việt / Latin cơ bản (Devanagari, CJK lạc, v.v.)
_NON_VIET_CHAR_RE = re.compile(r'[\u0900-\u097F\u4E00-\u9FFF\u3000-\u303F]+')


def _is_valid_viet_syllable(word: str) -> bool:
    """Kiểm tra xem từ có chứa tổ hợp nguyên âm hợp lệ hay không."""
    # Xử lý phụ âm ghép chứa nguyên âm: gi-, qu-
    # Trong "giấy", "gi" là phụ âm → 'i' không phải nguyên âm
    # Trong "quyền", "qu" là phụ âm → 'u' không phải nguyên âm
    w = word.lower()
    skip_indices = set()
    for i in range(len(w) - 1):
        if w[i] == 'g' and w[i + 1] == 'i' and i + 2 < len(w) and w[i + 2] in _STRIP_TONE:
            base_next = _STRIP_TONE.get(w[i + 2], w[i + 2]).lower()
            if base_next in ('a', 'ă', 'â', 'e', 'ê', 'o', 'ô', 'ơ', 'u', 'ư'):
                skip_indices.add(i + 1)  # skip 'i' in 'gi'
        if w[i] == 'q' and w[i + 1] == 'u':
            skip_indices.add(i + 1)  # skip 'u' in 'qu'

    # Trích xuất chuỗi nguyên âm liên tiếp
    vowel_runs = []
    current_run = []
    for idx, c in enumerate(word):
        if idx in skip_indices:
            # Ký tự này thuộc phụ âm ghép → xem như phụ âm
            if current_run:
                vowel_runs.append(current_run)
                current_run = []
            continue
        if c in _ALL_VOWELS:
            current_run.append(c)
        else:
            if current_run:
                vowel_runs.append(current_run)
                current_run = []
    if current_run:
        vowel_runs.append(current_run)

    for run in vowel_runs:
        if len(run) <= 1:
            continue

        # Strip dấu thanh và lowercase
        base_run = ''.join(_STRIP_TONE.get(c, c) for c in run).lower()

        # Đếm nguyên âm có dấu thanh trong run
        toned_in_run = sum(1 for c in run if c in _TONED_VOWELS)

        # Quy tắc 1: KHÔNG được có 2+ nguyên âm có dấu thanh liền nhau
        if toned_in_run >= 2:
            return False

        # Quy tắc 2: tổ hợp nguyên âm (base) phải nằm trong danh sách hợp lệ
        if base_run not in _VALID_VOWEL_CLUSTERS:
            # Thử cắt bớt: có thể run dài hơn 1 cluster (VD: 3+ nguyên âm)
            found_valid = False
            for length in (3, 2):
                if len(base_run) >= length:
                    prefix = base_run[:length]
                    suffix = base_run[length:]
                    if prefix in _VALID_VOWEL_CLUSTERS and (not suffix or len(suffix) <= 1):
                        found_valid = True
                        break
            if not found_valid and len(base_run) >= 2:
                # Tổ hợp nguyên âm không hợp lệ
                return False

    return True


def _remove_corrupted_words(text: str) -> str:
    """
    Phát hiện và xóa từ bị hỏng do AI rewriter.

    Từ hỏng = chứa tổ hợp nguyên âm không hợp lệ trong tiếng Việt.
    VD: "bấọng" (ấọ), "vòở" (òở), "chuyểình" (ểì), "khôơ" (ôơ),
        "củên" (ủê), "tạên" (ạê), "nóôi" (óôi), "nhìên" (ìê)

    Cũng xóa ký tự ngoại lai (Devanagari, CJK lạc, v.v.)
    """
    # 1. Xóa ký tự ngoại lai
    text = _NON_VIET_CHAR_RE.sub('', text)

    # 2. Xóa từ bị hỏng — tìm tất cả từ, kiểm tra từng từ
    def _check_word(match):
        word = match.group(0)
        # Từ quá ngắn (1-2 ký tự) → giữ
        if len(word) <= 2:
            return word
        # Từ chỉ chứa ASCII → có thể là tiếng Anh, để English map xử lý
        if word.isascii():
            return word
        # Kiểm tra tổ hợp nguyên âm
        if not _is_valid_viet_syllable(word):
            return ''
        return word

    # Tìm mọi "từ" chứa ký tự tiếng Việt
    text = re.sub(r'\b[\wÀ-ỹ]{3,}\b', _check_word, text, flags=re.UNICODE)

    # 3. Dọn dẹp: xóa dấu cách thừa sau khi xóa từ
    text = re.sub(r'  +', ' ', text)
    text = re.sub(r'(?m)^ +', '', text)

    return text


# Lỗi lặp từ (AI hay lặp 2-3 lần cùng cụm)
_REPEATED_WORD_RE = re.compile(
    r'\b([\wÀ-ỹ]{2,})\s+\1\b',
    re.IGNORECASE | re.UNICODE,
)

# Lỗi thiếu dấu cách sau dấu câu: "đi.Tôi" → "đi. Tôi"
_MISSING_SPACE_AFTER_PUNCT_RE = re.compile(
    r'([.!?,:;])([A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴĐa-zàáảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ])'
)

# Lỗi dấu câu kép: ".. " → ". ", "!! " → "! "
_DOUBLE_PUNCT_RE = re.compile(r'([.!?,;:])\1+')

# Dấu cách thừa trước dấu câu: "tôi ." → "tôi."
_SPACE_BEFORE_PUNCT_RE = re.compile(r'\s+([.!?,;:])')

# Lỗi dấu ngoặc kép không đóng/mở đúng
_QUOTE_FIXES = [
    ('"', '"'),  # smart quote mở
    ('"', '"'),  # smart quote đóng
    ('「', '"'),
    ('」', '"'),
    ('『', '"'),
    ('』', '"'),
]


def _fix_common_typos(text: str) -> str:
    """Sửa lỗi chính tả phổ biến từ AI rewrite."""

    # 0. Validator âm tiết tiếng Việt — sửa corruption ký tự tự động
    #    (thay cho việc liệt kê thủ công trong _OCR_MAP trước đây)
    text, stats = _vi_process_text(text)
    if stats.corrected:
        logger.info(
            "vi_validator: sửa %d/%d âm tiết hỏng (ratio=%.2f%%)",
            stats.corrected, stats.total_words, stats.invalid_ratio * 100,
        )
    if stats.uncorrectable:
        # Log tối đa 10 từ để tiện review nhưng tránh spam log
        sample = stats.uncorrectable[:10]
        logger.warning(
            "vi_validator: %d âm tiết không tự sửa được (ví dụ: %s)",
            len(stats.uncorrectable), sample,
        )

    # 0a. Xóa từ bị hỏng do AI rewriter (phần residual sau validator)
    text = _remove_corrupted_words(text)

    # 0b. Sửa "chít" → "chết" khi bị DeepSeek content filter thay thế
    #     Dùng helper giữ capitalization gốc (Cái chít → Cái chết)
    def _preserve_case(match, replacement):
        orig = match.group(0)
        if orig and orig[0].isupper():
            return replacement[0].upper() + replacement[1:]
        return replacement

    for pattern, replacement in _DEATH_EUPHEMISM_PATTERNS:
        if r'\1' in replacement:
            # Pattern có group backref → dùng re.sub bình thường
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        else:
            text = re.sub(
                pattern,
                lambda m, r=replacement: _preserve_case(m, r),
                text, flags=re.IGNORECASE,
            )

    # 1. Thêm dấu cách sau dấu câu bị thiếu: "đi.Tôi" → "đi. Tôi"
    text = _MISSING_SPACE_AFTER_PUNCT_RE.sub(r'\1 \2', text)

    # 2. Xóa dấu câu kép: ".." → ".", "!!" → "!"
    text = _DOUBLE_PUNCT_RE.sub(r'\1', text)

    # 2b. Sửa lỗi AI cắt giữa câu để lại dấu phẩy thừa:
    #     "gấp gáp,." → "gấp gáp."    (comma + period → period)
    #     "ngập ngừng,," → "ngập ngừng,"  (comma đôi)
    #     ",,, anh nói" → ", anh nói"
    text = re.sub(r',\s*\.', '.', text)
    text = re.sub(r',\s*,+', ',', text)

    # 2c. Xóa dấu phẩy LẺ LOI đứng một mình trên dòng do AI cắt mid-sentence
    #     Ví dụ: "Hơi thở ,\n" hoặc "\n ,\n" (dấu phẩy trước newline, không có từ tiếp)
    text = re.sub(r'\s+,(\s*\n)', r'\1', text)

    # 2d. Xóa dấu "," đầu dòng (artifact của truncation)
    text = re.sub(r'(\n)\s*,\s*', r'\1', text)

    # 3. Xóa dấu cách thừa trước dấu câu: "tôi ." → "tôi."
    text = _SPACE_BEFORE_PUNCT_RE.sub(r'\1', text)

    # 4. Sửa từ lặp liền nhau: "rằng rằng" → "rằng"
    text = _REPEATED_WORD_RE.sub(r'\1', text)

    # 5. Sửa cụm từ lặp phổ biến
    for wrong, right in _TYPO_MAP.items():
        if wrong != right and ' ' in wrong:
            # Chỉ sửa cụm từ (có dấu cách) — tránh sửa nhầm từ đơn hợp lệ
            text = re.sub(rf'\b{re.escape(wrong)}\b', right, text, flags=re.IGNORECASE)

    # 6. Chuẩn hóa dấu ngoặc kép
    for old, new in _QUOTE_FIXES:
        text = text.replace(old, new)

    # 7. Sửa dấu "..." bị thừa: "...." → "..."
    text = re.sub(r'\.{4,}', '...', text)

    # 8. Xóa từ ngập ngừng/suy nghĩ do AI thêm vào: "ưm...", "a...", "hừm...", "ừ..."
    #    VD: "chúng ta đã yêu nhau... a... ba năm" → "chúng ta đã yêu nhau... ba năm"
    text = re.sub(
        r'\.{2,3}\s*(?:ưm|uhm|à|a|ừ|hừm|hừ|ờ|ồ|ủa|eh|hmm|um|ah|oh|ơ|ê)\s*\.{2,3}',
        '...', text, flags=re.IGNORECASE,
    )
    # Dạng đứng đầu câu thoại: "Ưm... chúng ta" → "Chúng ta"
    text = re.sub(
        r'(?<=["\"])\s*(?:ưm|uhm|à|a|ừ|hừm|hừ|ờ|ồ|ủa|eh|hmm|um|ah|oh|ơ|ê)\s*\.{2,3}\s*',
        '', text, flags=re.IGNORECASE,
    )
    # Dạng có dấu phẩy: ", ưm," hoặc ", a,"
    text = re.sub(
        r',\s*(?:ưm|uhm|à|a|ừ|hừm|hừ|ờ|ồ|ủa|eh|hmm|um|ah|oh|ơ|ê)\s*,',
        ',', text, flags=re.IGNORECASE,
    )

    # 9. Sửa lỗi OCR / ký tự bị merge / hỏng
    def _ocr_replace(match, replacement):
        """Thay thế giữ nguyên viết hoa/thường của từ gốc."""
        orig = match.group(0)
        if orig[0].isupper():
            return replacement[0].upper() + replacement[1:]
        return replacement

    for wrong, right in _OCR_MAP.items():
        if ' ' in wrong:
            # Cụm có space → replace case-insensitive, giữ capitalization
            text = re.sub(
                re.escape(wrong),
                lambda m, r=right: _ocr_replace(m, r),
                text, flags=re.IGNORECASE,
            )
        else:
            text = re.sub(
                rf'\b{re.escape(wrong)}\b',
                lambda m, r=right: _ocr_replace(m, r),
                text, flags=re.IGNORECASE,
            )

    # 10. Sửa từ bị thiếu chữ đầu dòng (OCR cắt mất ký tự đầu)
    #     VD: "ìn về phía tôi" → "nhìn về phía tôi" (đầu câu)
    _HEAD_TRUNC = {
        'ìn':  'nhìn',   'ày': 'này',   'ôi': 'tôi',
        'ũng': 'cũng',   'ược': 'được', 'ưng': 'nhưng',
        'ười': 'người',   'ầy': 'đầy',  'ạy': 'vậy',
        'ồi': 'rồi',     'ại': 'lại',  'ốn': 'không',
        'ảy': 'vảy',     'ản': 'bản',  'ến': 'đến',
        'ật': 'thật',    'ức': 'lực',   'ổi': 'đổi',
    }
    for trunc, full in _HEAD_TRUNC.items():
        # Chỉ sửa khi đứng đầu dòng (sau \n hoặc đầu text)
        text = re.sub(
            rf'(?m)^{re.escape(trunc)}\b',
            full, text,
        )

    # 11. Thay từ tiếng Anh phổ biến → tiếng Việt
    #     Tránh tạo từ lặp: "giọng voice" → "giọng giọng" (sai) → chỉ giữ 1
    for eng, viet in _ENGLISH_VIET_MAP.items():
        def _eng_replace(m, r=viet):
            replacement = r[0].upper() + r[1:] if m.group(0)[0].isupper() else r
            # Kiểm tra từ trước/sau có trùng với replacement không
            start, end = m.start(), m.end()
            before = text[max(0, start - len(replacement) - 1):start].strip().lower()
            after = text[end:end + len(replacement) + 1].strip().lower()
            if before.endswith(replacement.lower()) or after.startswith(replacement.lower()):
                return ''  # xóa từ English, từ Việt đã có sẵn bên cạnh
            return replacement
        text = re.sub(rf'\b{re.escape(eng)}\b', _eng_replace, text, flags=re.IGNORECASE)
    # Dọn dẹp space thừa sau khi xóa English word
    text = re.sub(r'  +', ' ', text)

    # 12. Xóa dòng chỉ chứa dấu cách
    text = re.sub(r'(?m)^[ \t]+$', '', text)

    return text


def _add_para(doc, text: str, space_after_pt: float = 6.0):
    """Thêm 1 đoạn văn LEFT-align, space_after=space_after_pt pt."""
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(space_after_pt)
    return p


CHAR_LIMIT = 60_000   # ký tự tối đa mỗi file DOCX


def _write_txt(
    slug: str,
    part: int | None,
    title: str,
    author: str,
    chapters: list,
    output_dir: str,
    channel_name: str,
    total_parts: int,
) -> str:
    """Ghi một file TXT cho một nhóm chương. Trả về đường dẫn file."""
    suffix   = f"_p{part}" if part is not None else ""
    out_path = os.path.join(output_dir, f"{slug}{suffix}.txt")

    lines: list[str] = []

    # ── Lời chào ───────────────────────────────────────────────────────────
    if channel_name:
        if part is not None:
            greeting = (
                f"Chào mừng các bạn đến với {channel_name}. "
                f"Chúc các bạn nghe truyện vui vẻ. "
                f"(Phần {part}/{total_parts})"
            )
        else:
            greeting = (
                f"Chào mừng các bạn đến với {channel_name}. "
                f"Chúc các bạn nghe truyện vui vẻ."
            )
    else:
        greeting = f"{title} — {author}"
        if part is not None:
            greeting += f" (Phần {part}/{total_parts})"

    lines.append(greeting)
    lines.append('')

    # ── Nội dung ───────────────────────────────────────────────────────────
    for ch in chapters:
        content = clean_text(ch.get('content', ''))
        for line in content.splitlines():
            stripped = line.strip()
            if stripped:
                lines.append(stripped)
                lines.append('')

    # ── Lời kết ────────────────────────────────────────────────────────────
    if part is not None and part < total_parts:
        ending = (
            f"Hết phần {part}. Mời các bạn nghe tiếp phần {part + 1}."
            + (f" Nhớ like và đăng ký kênh {channel_name}." if channel_name else "")
        )
    else:
        ending = (
            f"Hoàn văn. Nhớ like và đăng ký kênh {channel_name} để nghe nhiều chuyện ngôn tình hay."
            if channel_name
            else "Hoàn văn."
        )
    lines.append(ending)

    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    return out_path


def _write_docx(
    slug: str,
    part: int | None,
    title: str,
    author: str,
    chapters: list,
    output_dir: str,
    channel_name: str,
    total_parts: int,
) -> str:
    """Ghi một file DOCX cho một nhóm chương. Trả về đường dẫn file."""
    from docx import Document
    from docx.shared import Pt

    suffix   = f"_p{part}" if part is not None else ""
    out_path = os.path.join(output_dir, f"{slug}{suffix}.docx")

    doc   = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # ── Lời chào ───────────────────────────────────────────────────────────
    doc.add_paragraph()

    if channel_name:
        if part is not None:
            greeting = (
                f"Chào mừng các bạn đến với {channel_name}. "
                f"Chúc các bạn nghe truyện vui vẻ. "
                f"(Phần {part}/{total_parts})"
            )
        else:
            greeting = (
                f"Chào mừng các bạn đến với {channel_name}. "
                f"Chúc các bạn nghe truyện vui vẻ."
            )
    else:
        greeting = f"{title} — {author}"
        if part is not None:
            greeting += f" (Phần {part}/{total_parts})"

    _add_para(doc, greeting)
    doc.add_paragraph()

    # ── Nội dung ───────────────────────────────────────────────────────────
    for ch in chapters:
        content = clean_text(ch.get('content', ''))
        for line in content.splitlines():
            stripped = line.strip()
            if stripped:
                _add_para(doc, stripped)
                doc.add_paragraph()

    # ── Lời kết ────────────────────────────────────────────────────────────
    if part is not None and part < total_parts:
        ending = (
            f"Hết phần {part}. Mời các bạn nghe tiếp phần {part + 1}."
            + (f" Nhớ like và đăng ký kênh {channel_name}." if channel_name else "")
        )
    else:
        ending = (
            f"Hoàn văn. Nhớ like và đăng ký kênh {channel_name} để nghe nhiều chuyện ngôn tình hay."
            if channel_name
            else "Hoàn văn."
        )
    _add_para(doc, ending)

    doc.save(out_path)
    return out_path


def stage_chapters(title: str, chapters: list, output_dir: str) -> str:
    """
    Lưu từng chương vào staging folder (JSON) để build DOCX sau.
    Trả về đường dẫn staging folder.

    Staging folder: <output_dir>/<slug>/chapters/<num:04d>.json
    """
    import json

    slug        = _slugify(title)
    stage_dir   = os.path.join(output_dir, slug, 'chapters')
    os.makedirs(stage_dir, exist_ok=True)

    for ch in chapters:
        num      = ch.get('number', 0)
        filename = os.path.join(stage_dir, f"{num:04d}.json")
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(ch, f, ensure_ascii=False, indent=2)

    logger.info(f"  📂 Staged {len(chapters)} chương → {stage_dir}")
    return stage_dir


def load_staged_chapters(title: str, output_dir: str) -> list:
    """
    Đọc tất cả chương đã staged, sort theo số chương.
    Trả về list chapters đã sort.
    """
    import json

    slug      = _slugify(title)
    stage_dir = os.path.join(output_dir, slug, 'chapters')

    if not os.path.isdir(stage_dir):
        return []

    chapters = []
    for fname in sorted(os.listdir(stage_dir)):
        if not fname.endswith('.json'):
            continue
        fpath = os.path.join(stage_dir, fname)
        try:
            with open(fpath, 'r', encoding='utf-8') as f:
                chapters.append(json.load(f))
        except Exception as e:
            logger.warning(f"  ⚠️ Không đọc được {fname}: {e}")

    # Sort theo số chương (đề phòng tên file bị lệch)
    chapters.sort(key=lambda c: c.get('number', 0))
    return chapters


def build_docx_from_staged(
    title: str,
    author: str,
    output_dir: str,
    channel_name: str = '',
) -> list[str] | None:
    """
    Đọc tất cả staged chapters, sort theo số chương, rồi build DOCX.
    Xóa các file DOCX cũ của truyện trước khi build lại.
    Trả về list đường dẫn file DOCX đã tạo.
    """
    chapters = load_staged_chapters(title, output_dir)
    if not chapters:
        logger.warning(f"  ⚠️ Không có chương nào trong staging cho '{title}'")
        return None

    logger.info(f"  🔨 Build DOCX từ {len(chapters)} chương đã staged (sort theo số chương)...")
    return _build_docx(title, author, chapters, output_dir, channel_name)


def _build_docx(
    title: str,
    author: str,
    chapters: list,
    output_dir: str,
    channel_name: str = '',
) -> list[str] | None:
    """
    Build DOCX từ danh sách chương (đã được sort trước).
    Tự động tách file nếu vượt CHAR_LIMIT.
    Xóa các file DOCX cũ trước khi ghi mới.
    """
    try:
        import docx as _docx_check  # noqa: F401
    except ImportError:
        logger.error("❌ Thiếu thư viện python-docx. Chạy: pip install python-docx")
        return None

    slug       = _slugify(title)
    novel_dir  = os.path.join(output_dir, slug)
    os.makedirs(novel_dir, exist_ok=True)

    # Xóa các file DOCX/TXT cũ để tránh phần thừa sau khi rebuild
    # Dùng try/except để không crash nếu file đang mở hoặc bị khóa
    for f in os.listdir(novel_dir):
        if (f.endswith('.docx') or f.endswith('.txt')) and not f.startswith('~$'):
            try:
                os.remove(os.path.join(novel_dir, f))
            except OSError as e:
                logger.warning(f"  ⚠️ Không thể xóa file cũ {f}: {e}")

    total_chars = sum(len(ch.get('content', '')) for ch in chapters)

    if total_chars <= CHAR_LIMIT:
        out_path = _write_docx(slug, None, title, author, chapters, novel_dir, channel_name, 1)
        txt_path = _write_txt(slug, None, title, author, chapters, novel_dir, channel_name, 1)
        logger.info(f"  📄 DOCX: {out_path}  ({len(chapters)} chương, {total_chars:,} ký tự)")
        logger.info(f"  📝 TXT:  {txt_path}")
        return [out_path]

    # Tách thành nhiều phần theo CHAR_LIMIT
    groups: list[list] = []
    current_group: list = []
    current_chars = 0

    for ch in chapters:
        ch_len = len(ch.get('content', ''))
        if current_group and current_chars + ch_len > CHAR_LIMIT:
            groups.append(current_group)
            current_group = []
            current_chars = 0
        current_group.append(ch)
        current_chars += ch_len
    if current_group:
        groups.append(current_group)

    total_parts = len(groups)
    out_paths: list[str] = []

    for idx, group in enumerate(groups, start=1):
        out_path = _write_docx(slug, idx, title, author, group, novel_dir, channel_name, total_parts)
        txt_path = _write_txt(slug, idx, title, author, group, novel_dir, channel_name, total_parts)
        group_chars = sum(len(ch.get('content', '')) for ch in group)
        logger.info(
            f"  📄 DOCX phần {idx}/{total_parts}: {out_path}"
            f"  ({len(group)} chương, {group_chars:,} ký tự)"
        )
        logger.info(f"  📝 TXT  phần {idx}/{total_parts}: {txt_path}")
        out_paths.append(out_path)

    return out_paths


def save_novel_as_docx(
    title: str,
    author: str,
    chapters: list,
    output_dir: str,
    channel_name: str = '',
) -> list[str] | None:
    """
    Stage các chương mới rồi build lại toàn bộ DOCX từ staged chapters.
    Đảm bảo DOCX luôn được sort theo số chương và đúng giới hạn ký tự.
    """
    # 1. Stage các chương mới (ghi đè nếu đã có)
    stage_chapters(title, chapters, output_dir)

    # 2. Load tất cả staged chapters (bao gồm cả cũ + mới), đã sort
    all_chapters = load_staged_chapters(title, output_dir)

    # 3. Build lại toàn bộ DOCX
    return _build_docx(title, author, all_chapters, output_dir, channel_name)
